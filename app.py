"""
Flask Web Application for DOCX Translation
English to Telugu Translation Service
"""

from flask import Flask, render_template, request, send_file, jsonify, flash, session, Response, stream_with_context
from werkzeug.utils import secure_filename
import os
import tempfile
from pathlib import Path
from docx import Document
import traceback
import json
import uuid
from datetime import datetime

# Import new modules
from translate_docx_unified import UnifiedDOCXTranslator
from translation_engine import get_available_engines
from auth import authenticate, get_user_id
from feedback_logging import log_translation, log_feedback

try:
    from supabase_config import save_feedback
    HAS_SUPABASE = True
except ImportError:
    HAS_SUPABASE = False
    # Silent - Supabase is optional and only needed for comments/feedback features

try:
    from flask_session import Session
    HAS_FLASK_SESSION = True
except ImportError:
    HAS_FLASK_SESSION = False
    print("Warning: flask-session not installed. Using default session.")

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv("FLASK_SECRET_KEY", "your-secret-key-here-change-in-production")
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv("FLASK_MAX_CONTENT_LENGTH", 50 * 1024 * 1024))  # 50MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
if HAS_FLASK_SESSION:
    app.config['SESSION_TYPE'] = 'filesystem'
    Session(app)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Store file mappings (in production, use Redis or database)
file_mappings = {}

# Store translation metadata
translation_metadata = {}  # translation_id -> {user_id, engine, etc.}

# Pre-load models on startup
def preload_models():
    """Pre-load all translation models"""
    print("=" * 60)
    print("Pre-loading translation models...")
    print("=" * 60)
    
    available_engines = get_available_engines()
    
    # Load models in order: Gemini, Google Standard, Google Adaptive, IndicTrans2
    load_order = ['gemini-3-pro', 'google-standard', 'google-adaptive', 'indictrans2']
    
    for engine in load_order:
        if engine in available_engines and available_engines[engine]['available']:
            try:
                print(f"Loading {engine}...")
                translator = UnifiedDOCXTranslator(engine=engine)
                # Just initialize, don't translate
                print(f"✓ {engine} ready")
            except Exception as e:
                print(f"✗ {engine} failed to load: {e}")
    
    print("=" * 60)
    print("Model pre-loading complete!")
    print("=" * 60)

# Pre-load models on import (runs when app starts)
try:
    preload_models()
except Exception as e:
    print(f"Warning: Error pre-loading models: {e}")

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/login', methods=['POST'])
def login():
    """Authenticate user"""
    try:
        data = request.get_json()
        username = data.get('username', '')
        password = data.get('password', '')
        
        if not username or not password:
            return jsonify({'error': 'Username and password required'}), 400
        
        success, user_id, error = authenticate(username, password)
        
        if success:
            session['user_id'] = user_id
            session['username'] = username
            return jsonify({'success': True, 'user_id': user_id, 'username': username})
        else:
            return jsonify({'error': error or 'Authentication failed'}), 401
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/logout', methods=['POST'])
def logout():
    """Logout user"""
    session.pop('user_id', None)
    session.pop('username', None)
    return jsonify({'success': True})

@app.route('/engines', methods=['GET'])
def get_engines():
    """Get available translation engines"""
    engines = get_available_engines()
    return jsonify({'success': True, 'engines': engines})

def extract_text_from_docx(file_path, as_html=False):
    """Extract text content from DOCX file, optionally as HTML (preserves tables)"""
    try:
        doc = Document(file_path)
        if as_html:
            # Extract as HTML to preserve tables and formatting
            html_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    html_parts.append(f"<p>{para.text}</p>")
            
            # Extract tables as HTML tables
            for table in doc.tables:
                html_parts.append("<table style='border-collapse: collapse; width: 100%; margin: 15px 0;'>")
                for row_idx, row in enumerate(table.rows):
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                        # First row as header
                        tag = "th" if row_idx == 0 else "td"
                        html_parts.append(f"<{tag} style='border: 1px solid #ddd; padding: 8px; text-align: left;'>{cell_text}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")
            
            return '\n'.join(html_parts)
        else:
            # Plain text extraction (for translation)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                        if cell_text.strip():
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraphs.append(' | '.join(row_texts))
            return '\n\n'.join(paragraphs)
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

@app.route('/translate', methods=['POST'])
def translate_file():
    """Handle file upload and generate all 4 translations simultaneously"""
    try:
        # Check authentication
        user_id = session.get('user_id') or request.form.get('user_id')
        if not user_id:
            return jsonify({'error': 'Authentication required. Please login first.'}), 401
        
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        # Check if file is selected
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check if file type is allowed
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload a .docx file'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        
        try:
            # Extract original text for preview
            original_text = extract_text_from_docx(input_path, as_html=True)  # Get HTML for display
            original_text_plain = extract_text_from_docx(input_path, as_html=False)  # Plain text for translation
            
            # Generate translation ID
            translation_id = str(uuid.uuid4())
            input_file = Path(input_path)
            
            # Get available engines
            available_engines = get_available_engines()
            # Order: Gemini first (fastest), then Google Standard, Google Adaptive, IndicTrans2 last (slowest)
            # Always include IndicTrans2 - it will handle errors gracefully
            engines_to_use = ['gemini-3-pro', 'google-standard', 'google-adaptive', 'indictrans2']
            
            # Generate all translations (will be updated progressively)
            translations = {}
            translation_files = {}
            
            # Start translations in parallel but in order
            import asyncio
            import concurrent.futures
            
            def translate_with_engine(engine):
                """Translate with a specific engine"""
                # Check if engine is available
                if engine not in available_engines or not available_engines[engine]['available']:
                    error_msg = f'Engine {engine} is not available. Please install required dependencies.'
                    print(f"Warning: {error_msg}")
                    return engine, {
                        'text': '',
                        'error': error_msg,
                        'filename': None
                    }
                
                try:
                    # Create translator for this engine
                    translator = UnifiedDOCXTranslator(engine=engine)
                    
                    # Create output filename
                    output_filename = f"{input_file.stem}_telugu_{engine}{input_file.suffix}"
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    
                    # Translate the document
                    translator.translate_docx(
                        input_path, 
                        output_path,
                        user_id=user_id,
                        translation_id=f"{translation_id}_{engine}"
                    )
                    
                    # Extract translated text as HTML for display
                    translated_text = extract_text_from_docx(output_path, as_html=True)
                    
                    result = {
                        'text': translated_text,
                        'filename': output_filename,
                        'error': None
                    }
                    
                    translation_files[engine] = {
                        'path': output_path,
                        'filename': output_filename
                    }
                    
                    # Log translation
                    log_translation(user_id, engine, original_text_plain, extract_text_from_docx(output_path, as_html=False), translation_id)
                    
                    return engine, result
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error translating with {engine}: {error_msg}")
                    return engine, {
                        'text': '',
                        'error': error_msg,
                        'filename': None
                    }
            
            # Store initial state for progressive updates
            file_mappings[translation_id] = {
                'input_path': input_path,
                'original_text': original_text,
                'translations': {},
                'status': {engine: 'pending' for engine in engines_to_use}
            }
            
            translation_metadata[translation_id] = {
                'user_id': user_id,
                'engines': engines_to_use,
                'created_at': datetime.now().isoformat()
            }
            
            # Start background translation with PARALLEL execution
            import threading
            import concurrent.futures
            
            def translate_single_engine(engine):
                """Translate with a single engine"""
                if translation_id not in file_mappings:
                    return
                
                # Update status to translating
                file_mappings[translation_id]['status'][engine] = 'translating'
                
                # Double-check engine is available (shouldn't happen if filtering worked)
                if engine not in available_engines or not available_engines[engine]['available']:
                    if engine == 'indictrans2':
                        error_msg = "IndicTrans2 is not available. Please check the API connection."
                    else:
                        error_msg = f"Engine {engine} is not available. Please install required dependencies."
                    print(f"Warning: {error_msg}")
                    file_mappings[translation_id]['status'][engine] = 'error'
                    file_mappings[translation_id]['translations'][engine] = {
                        'text': '',
                        'error': error_msg,
                        'filename': None
                    }
                    return
                
                try:
                    # Create translator for this engine
                    translator = UnifiedDOCXTranslator(engine=engine)
                    
                    # Create output filename
                    output_filename = f"{input_file.stem}_telugu_{engine}{input_file.suffix}"
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    
                    # Translate the document
                    translator.translate_docx(
                        input_path, 
                        output_path,
                        user_id=user_id,
                        translation_id=f"{translation_id}_{engine}"
                    )
                    
                    # Extract translated text
                    try:
                        translated_text = extract_text_from_docx(output_path, as_html=True)
                    except Exception as extract_error:
                        print(f"Warning: Could not extract text from {output_path}: {extract_error}")
                        translated_text = ""  # Use empty text if extraction fails
                    
                    # Store in file_mappings
                    file_mappings[translation_id]['translations'][engine] = {
                        'path': output_path,
                        'filename': output_filename,
                        'text': translated_text
                    }
                    
                    # Update status - IMPORTANT: Set status AFTER storing translations
                    # Use direct assignment to ensure it's set immediately
                    if translation_id in file_mappings:
                        file_mappings[translation_id]['status'][engine] = 'completed'
                        
                        # Verify it was set correctly
                        actual_status = file_mappings[translation_id]['status'].get(engine)
                        if actual_status != 'completed':
                            print(f"[ERROR] Status not set correctly! Expected 'completed', got '{actual_status}'")
                        else:
                            print(f"[STATUS UPDATE] ✓ {engine}: status set to 'completed' for translation_id={translation_id}")
                            print(f"[STATUS DEBUG] Full status dict: {dict(file_mappings[translation_id]['status'])}")
                            print(f"[STATUS DEBUG] Has translations entry: {engine in file_mappings[translation_id].get('translations', {})}")
                    else:
                        print(f"[ERROR] translation_id {translation_id} not found in file_mappings when trying to set status!")
                    
                    # Log translation
                    log_translation(user_id, engine, original_text_plain, extract_text_from_docx(output_path, as_html=False), translation_id)
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error translating with {engine}: {error_msg}")
                    file_mappings[translation_id]['status'][engine] = 'error'
                    file_mappings[translation_id]['translations'][engine] = {
                        'text': '',
                        'error': error_msg,
                        'filename': None
                    }
            
            def translate_all_parallel():
                """Translate with all engines in PARALLEL using ThreadPoolExecutor"""
                # Use ThreadPoolExecutor to run all engines in parallel
                # No timeout - let IndicTrans2 take as long as it needs
                with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
                    # Submit all translation tasks in parallel
                    futures = {executor.submit(translate_single_engine, engine): engine 
                              for engine in engines_to_use}
                    
                    # Wait for all to complete (they run in parallel) - no timeout
                    for future in concurrent.futures.as_completed(futures):
                        engine = futures[future]
                        try:
                            future.result()  # Get result (or raise exception)
                        except Exception as e:
                            print(f"Translation task for {engine} raised exception: {e}")
            
            # Start parallel translation in background thread
            thread = threading.Thread(target=translate_all_parallel)
            thread.daemon = True
            thread.start()
            
            # Return immediately with original text and translation ID
            # Frontend will poll for status updates
            return jsonify({
                'success': True,
                'original_text': original_text,
                'translation_id': translation_id,
                'engines': engines_to_use,
                'status': 'started'
            })
        
        except Exception as e:
            error_msg = str(e)
            print(f"Error: {error_msg}")
            print(traceback.format_exc())
            return jsonify({'error': f'Translation failed: {error_msg}'}), 500
    
    except Exception as e:
        error_msg = str(e)
        print(f"Error: {error_msg}")
        print(traceback.format_exc())
        return jsonify({'error': f'Translation failed: {error_msg}'}), 500

@app.route('/translate-status/<translation_id>', methods=['GET'])
def get_translation_status(translation_id):
    """Get status of translations - returns completed translations"""
    try:
        if translation_id not in file_mappings:
            return jsonify({'error': 'Translation session not found'}), 404
        
        file_info = file_mappings[translation_id]
        translations = {}
        status = file_info.get('status', {})
        
        # Debug: Print what we're returning
        print(f"[STATUS ENDPOINT] translation_id={translation_id}, status dict: {status}")
        
        # Build translations object with current status
        # IMPORTANT: Check status first, then translations
        for engine in file_info.get('status', {}).keys():
            current_status = status.get(engine, 'pending')
            print(f"[STATUS ENDPOINT] {engine}: status={current_status}, has_translation={engine in file_info.get('translations', {})}")
            
            if engine in file_info.get('translations', {}):
                trans_data = file_info['translations'][engine]
                translations[engine] = {
                    'text': trans_data.get('text', ''),
                    'filename': trans_data.get('filename'),
                    'error': trans_data.get('error'),
                    'status': current_status  # Use current_status from status dict
                }
            else:
                # Still translating or pending
                translations[engine] = {
                    'text': '',
                    'filename': None,
                    'error': None,
                    'status': current_status  # Use current_status from status dict
                }
        
        # Check if all are complete
        all_complete = all(s in ['completed', 'error'] for s in status.values())
        
        response_data = {
            'success': True,
            'translations': translations,
            'status': status,
            'all_complete': all_complete
        }
        
        # Debug: Print what we're sending (without full text to avoid clutter)
        debug_data = {
            'success': response_data['success'],
            'status': response_data['status'],
            'all_complete': response_data['all_complete'],
            'translations': {k: {'has_text': bool(v.get('text')), 'status': v.get('status'), 'filename': v.get('filename')} for k, v in response_data['translations'].items()}
        }
        print(f"[STATUS ENDPOINT] Returning response for {translation_id}:")
        print(json.dumps(debug_data, indent=2))
        
        return jsonify(response_data)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download', methods=['POST'])
def download_file():
    """Download a specific translated file"""
    try:
        data = request.get_json()
        translation_id = data.get('translation_id')
        engine = data.get('engine')
        
        if not translation_id or translation_id not in file_mappings:
            return jsonify({'error': 'Translation session not found'}), 404
        
        file_info = file_mappings[translation_id]
        
        if engine:
            # Download specific engine translation
            if 'translations' in file_info and engine in file_info['translations']:
                output_path = file_info['translations'][engine]['path']
                output_filename = file_info['translations'][engine]['filename']
            else:
                return jsonify({'error': 'Translation file not found'}), 404
        else:
            # Download first available translation
            if 'translations' in file_info and file_info['translations']:
                first_engine = list(file_info['translations'].keys())[0]
                output_path = file_info['translations'][first_engine]['path']
                output_filename = file_info['translations'][first_engine]['filename']
            else:
                return jsonify({'error': 'No translation files found'}), 404
        
        if not os.path.exists(output_path):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete-translation', methods=['POST'])
def delete_translation():
    """Delete a specific translation file"""
    try:
        data = request.get_json()
        translation_id = data.get('translation_id')
        engine = data.get('engine')
        
        if not translation_id or translation_id not in file_mappings:
            return jsonify({'error': 'Translation session not found'}), 404
        
        if not engine:
            return jsonify({'error': 'Engine name required'}), 400
        
        file_info = file_mappings[translation_id]
        
        if 'translations' in file_info and engine in file_info['translations']:
            file_path = file_info['translations'][engine]['path']
            
            # Delete file
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Deleted translation file: {file_path}")
            
            # Remove from mappings
            del file_info['translations'][engine]
            
            return jsonify({'success': True, 'message': f'Translation {engine} deleted'})
        else:
            return jsonify({'error': 'Translation file not found'}), 404
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/comments', methods=['POST'])
def create_comment():
    """Create a new comment"""
    try:
        data = request.get_json()
        
        # Get user_id from session if available
        user_id = session.get('user_id')
        
        comment_data = {
            'translation_id': data.get('translation_id'),
            'doc_type': data.get('doc_type'),
            'engine': data.get('engine'),  # Which translation engine this comment is for
            'selected_text': data.get('selected_text', ''),
            'comment': data.get('comment', ''),
            'thumbs_rating': data.get('thumbs_rating')  # thumbs up/down rating
            # Note: text_position removed as it was causing Supabase errors
        }
        
        # Add user_id if available
        if user_id:
            comment_data['user_id'] = user_id
        
        if not comment_data['translation_id'] or not comment_data['doc_type'] or not comment_data['comment']:
            return jsonify({'error': 'Missing required fields'}), 400
        
        if HAS_SUPABASE:
            from supabase_config import save_comment
            result = save_comment(comment_data)
            if result.get('success'):
                return jsonify({'success': True, 'comment': result.get('data')})
            else:
                return jsonify({'success': False, 'error': result.get('error')}), 500
        else:
            # Supabase is required for comments - inform user
            return jsonify({
                'success': False, 
                'error': 'Comments feature requires Supabase. Please install: pip install supabase'
            }), 503  # 503 Service Unavailable
    
    except Exception as e:
        print(f"Error in create_comment: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/comments', methods=['GET'])
def get_comments():
    """Get comments for a translation"""
    try:
        translation_id = request.args.get('translation_id')
        engine = request.args.get('engine')  # Optional: filter by engine
        
        if not translation_id:
            return jsonify({'error': 'translation_id required'}), 400
        
        try:
            if HAS_SUPABASE:
                from supabase_config import get_comments
                result = get_comments(translation_id, engine=engine)
                if result.get('success'):
                    return jsonify({'success': True, 'comments': result.get('data', [])})
                else:
                    # If Supabase fails, return empty comments silently
                    # Only log if it's a real error (not just "not configured" or table missing)
                    error_str = str(result.get('error', ''))
                    if ("not configured" not in error_str.lower() and 
                        "PGRST205" not in error_str and
                        "Could not find the table" not in error_str and
                        "schema cache" not in error_str.lower()):
                        print(f"Warning: Failed to get comments from Supabase: {result.get('error')}")
                    return jsonify({'success': True, 'comments': []})
            else:
                # Supabase not installed - return empty comments silently
                # User will only see this when they try to comment/highlight
                return jsonify({'success': True, 'comments': []})
        except Exception as supabase_error:
            # If Supabase call fails, return empty comments silently
            return jsonify({'success': True, 'comments': []})
    
    except Exception as e:
        print(f"Error in get_comments endpoint: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/comments/<int:comment_id>', methods=['DELETE'])
def delete_comment(comment_id):
    """Delete a comment"""
    try:
        if HAS_SUPABASE:
            from supabase_config import delete_comment
            result = delete_comment(comment_id)
            if result.get('success'):
                return jsonify({'success': True})
            else:
                return jsonify({'success': False, 'error': result.get('error')}), 500
        else:
            return jsonify({'success': True})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save-edited', methods=['POST'])
def save_edited_translation():
    """Save edited translation back to DOCX"""
    try:
        data = request.get_json()
        edited_text = data.get('text', '')
        translation_id = data.get('translation_id')
        
        if not edited_text:
            return jsonify({'error': 'No text provided'}), 400
        
        if not translation_id or translation_id not in file_mappings:
            return jsonify({'error': 'Translation session not found'}), 400
        
        output_path = file_mappings[translation_id]['output_path']
        
        # Update the DOCX file with edited text
        doc = Document(output_path)
        paragraphs = edited_text.split('\n\n')
        
        # Clear existing paragraphs (except first which might be title)
        for para in doc.paragraphs:
            para.clear()
        
        # Add edited paragraphs
        for i, para_text in enumerate(paragraphs):
            if para_text.strip():
                if i < len(doc.paragraphs):
                    doc.paragraphs[i].text = para_text
                else:
                    doc.add_paragraph(para_text)
        
        doc.save(output_path)
        
        return jsonify({'success': True, 'message': 'Document updated successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/feedback', methods=['POST'])
def submit_feedback():
    """Submit user feedback - accepts partial updates"""
    try:
        data = request.get_json()
        
        # Get user_id and translation_id
        user_id = session.get('user_id') or data.get('user_id')
        translation_id = data.get('translation_id')
        
        # Get translation model from metadata
        translation_model = 'indictrans2'  # default
        if translation_id and translation_id in translation_metadata:
            translation_model = translation_metadata[translation_id].get('engine', 'indictrans2')
        else:
            translation_model = data.get('translation_model', 'indictrans2')
        
        # Build feedback data with defaults, allowing partial updates
        criteria_ratings = data.get('criteria_ratings', {})
        
        feedback_data = {
            'file_type': data.get('file_type', ''),
            'language_pair': data.get('language_pair', 'English-Telugu'),
            'translation_method': data.get('translation_method', translation_model),
            'overall_quality': data.get('overall_quality'),
            'structure_preservation': data.get('structure_preservation'),
            'preview_features': data.get('preview_features'),
            'suggestions': data.get('suggestions', ''),
            'thumbs_rating': data.get('thumbs_rating'),  # Add thumbs rating support
            'criteria_ratings': criteria_ratings if criteria_ratings else None,  # Store criteria ratings as JSON
            'created_at': datetime.now().isoformat()
        }
        
        # Only submit if at least one field has a value
        if any([
            feedback_data['file_type'],
            feedback_data['overall_quality'] is not None,
            feedback_data['structure_preservation'] is not None,
            feedback_data['preview_features'] is not None,
            feedback_data['suggestions'],
            feedback_data['thumbs_rating'],  # Include thumbs rating
            criteria_ratings  # Include criteria ratings
        ]):
            # Log feedback with user_id and translation_model
            if user_id:
                result = log_feedback(user_id, translation_id, translation_model, feedback_data)
                if result.get('success'):
                    return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
                else:
                    # Supabase is required for feedback - inform user
                    error_msg = result.get('error', 'Failed to save feedback')
                    if 'not configured' in error_msg.lower() or 'not installed' in error_msg.lower():
                        return jsonify({
                            'success': False, 
                            'error': 'Feedback feature requires Supabase. Please install: pip install supabase'
                        }), 503
                    return jsonify({'success': False, 'error': error_msg}), 500
            elif HAS_SUPABASE:
                # Fallback to old method if no user_id
                result = save_feedback(feedback_data)
                if result.get('success'):
                    return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
                else:
                    return jsonify({'success': False, 'error': result.get('error', 'Failed to save feedback')}), 500
            else:
                # Supabase is required for feedback - inform user
                return jsonify({
                    'success': False, 
                    'error': 'Feedback feature requires Supabase. Please install: pip install supabase'
                }), 503
        else:
            return jsonify({'success': False, 'error': 'No feedback data provided'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/feedback', methods=['GET'])
def get_feedback():
    """Get feedback for a translation"""
    try:
        translation_id = request.args.get('translation_id')
        engine = request.args.get('engine')
        user_id = session.get('user_id') or request.args.get('user_id')
        
        if not translation_id:
            return jsonify({'error': 'translation_id required'}), 400
        
        if HAS_SUPABASE:
            from supabase_config import get_supabase_client
            supabase = get_supabase_client()
            if supabase:
                query = supabase.table("feedback").select("*").eq("translation_id", translation_id)
                
                if engine:
                    query = query.eq("translation_method", engine)
                if user_id:
                    query = query.eq("user_id", user_id)
                
                response = query.order("created_at", desc=True).limit(1).execute()
                
                if response.data and len(response.data) > 0:
                    feedback = response.data[0]
                    # Check if it has criteria_ratings (review submitted)
                    has_criteria = feedback.get('criteria_ratings') is not None
                    return jsonify({
                        'success': True,
                        'feedback': feedback,
                        'has_review': has_criteria
                    })
        
        return jsonify({'success': True, 'feedback': None, 'has_review': False})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/translation-history', methods=['GET'])
def get_translation_history():
    """Get all previous translations for the current user"""
    try:
        user_id = session.get('user_id') or request.args.get('user_id')
        if not user_id:
            print("[HISTORY] No user_id found in session or request")
            return jsonify({'error': 'Authentication required'}), 401
        
        print(f"[HISTORY] Fetching history for user_id: {user_id}")
        
        if HAS_SUPABASE:
            from supabase_config import get_supabase_client
            supabase = get_supabase_client()
            if supabase:
                # Get all translations for this user, grouped by translation_id
                print(f"[HISTORY] Querying translation_logs for user_id: {user_id}")
                translations_response = supabase.table("translation_logs").select("*").eq("user_id", user_id).order("created_at", desc=True).execute()
                
                print(f"[HISTORY] Found {len(translations_response.data)} translation log entries")
                
                # Group translations by translation_id
                translation_sessions = {}
                translations_without_id = []
                for trans in translations_response.data:
                    trans_id = trans.get('translation_id')
                    # Handle None, empty string, or missing translation_id
                    if not trans_id or trans_id.strip() == '':
                        translations_without_id.append(trans.get('id'))
                        print(f"[HISTORY] Skipping translation without translation_id: id={trans.get('id')}, model={trans.get('translation_model')}")
                        continue
                    
                    trans_id = trans_id.strip()  # Clean whitespace
                    
                    if trans_id not in translation_sessions:
                        translation_sessions[trans_id] = {
                            'translation_id': trans_id,
                            'created_at': trans.get('created_at'),
                            'translations': {},
                            'feedback': {},
                            'comments': []
                        }
                    
                    engine = trans.get('translation_model')
                    translation_sessions[trans_id]['translations'][engine] = {
                        'source_text': trans.get('source_text', ''),
                        'translated_text': trans.get('translated_text', ''),
                        'engine': engine,
                        'created_at': trans.get('created_at')
                    }
                
                print(f"[HISTORY] Grouped into {len(translation_sessions)} translation sessions")
                if translations_without_id:
                    print(f"[HISTORY] Warning: {len(translations_without_id)} translations were skipped due to missing translation_id")
                
                # Get feedback for each translation_id
                for trans_id in translation_sessions.keys():
                    try:
                        feedback_response = supabase.table("feedback").select("*").eq("translation_id", trans_id).eq("user_id", user_id).order("created_at", desc=True).execute()
                        print(f"[HISTORY] Found {len(feedback_response.data)} feedback entries for translation_id: {trans_id}")
                        
                        # Use a dict to track the most recent feedback per engine
                        engine_feedback_map = {}
                        for feedback in feedback_response.data:
                            # Prioritize translation_method as it has the correct engine name
                            engine = feedback.get('translation_method') or feedback.get('translation_model')
                            print(f"[HISTORY] Processing feedback entry - translation_model: {feedback.get('translation_model')}, translation_method: {feedback.get('translation_method')}, engine: {engine}, has_criteria_ratings: {bool(feedback.get('criteria_ratings'))}")
                            if engine:
                                # Only keep the first (most recent) feedback for each engine
                                # Since we order by created_at desc, first occurrence is most recent
                                if engine not in engine_feedback_map:
                                    engine_feedback_map[engine] = {
                                        'overall_quality': feedback.get('overall_quality'),
                                        'structure_preservation': feedback.get('structure_preservation'),
                                        'preview_features': feedback.get('preview_features'),
                                        'thumbs_rating': feedback.get('thumbs_rating'),
                                        'criteria_ratings': feedback.get('criteria_ratings'),
                                        'suggestions': feedback.get('suggestions'),
                                        'created_at': feedback.get('created_at')
                                    }
                                    print(f"[HISTORY] Added feedback for engine: {engine}")
                        
                        print(f"[HISTORY] Collected feedback for {len(engine_feedback_map)} engines: {list(engine_feedback_map.keys())}")
                        
                        # Add all collected feedback to the session
                        for engine, feedback_data in engine_feedback_map.items():
                            translation_sessions[trans_id]['feedback'][engine] = feedback_data
                    except Exception as e:
                        print(f"[HISTORY] Error fetching feedback for {trans_id}: {e}")
                    
                    # Get comments for this translation
                    try:
                        comments_response = supabase.table("comments").select("*").eq("translation_id", trans_id).order("created_at", desc=True).execute()
                        translation_sessions[trans_id]['comments'] = comments_response.data if comments_response.data else []
                    except Exception as e:
                        print(f"[HISTORY] Error fetching comments for {trans_id}: {e}")
                        translation_sessions[trans_id]['comments'] = []
                
                # Convert to list sorted by date
                history_list = list(translation_sessions.values())
                history_list.sort(key=lambda x: x.get('created_at', ''), reverse=True)
                
                print(f"[HISTORY] Returning {len(history_list)} history sessions")
                return jsonify({
                    'success': True,
                    'history': history_list
                })
            else:
                print("[HISTORY] Supabase client is None")
                return jsonify({'success': True, 'history': [], 'message': 'Supabase not configured'})
        else:
            print("[HISTORY] HAS_SUPABASE is False")
            # If Supabase not available, return empty history
            return jsonify({'success': True, 'history': [], 'message': 'Supabase not available'})
    
    except Exception as e:
        print(f"[HISTORY] Error getting translation history: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'service': 'DOCX Translation Service'})

if __name__ == '__main__':
    # Print startup message
    print("=" * 60)
    print("DOCX Translation Web Service")
    print("=" * 60)
    print("Starting server...")
    print("Access the web interface at: http://localhost:5000")
    print("=" * 60)
    
    app.run(debug=True, host='0.0.0.0', port=5000)

