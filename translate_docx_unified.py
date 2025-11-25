"""
Unified DOCX Translator
Supports multiple translation engines
"""

from docx import Document
import os
from pathlib import Path
from translation_engine import translate_batch, get_available_engines
from glossary import get_glossary
from feedback_logging import log_translation

class UnifiedDOCXTranslator:
    """Unified DOCX Translator supporting multiple engines"""
    
    def __init__(self, engine="indictrans2"):
        """
        Initialize translator
        
        Args:
            engine: Translation engine to use
        """
        self.engine = engine
        self.glossary = get_glossary()
        print(f"Initialized UnifiedDOCXTranslator with engine: {engine}")
    
    def translate_docx(self, input_path, output_path=None, user_id=None, translation_id=None):
        """
        Translate a DOCX file
        
        Args:
            input_path: Path to input DOCX file
            output_path: Path to output DOCX file (optional)
            user_id: User ID for logging
            translation_id: Translation session ID for logging
            
        Returns:
            Path to translated file
        """
        if output_path is None:
            input_file = Path(input_path)
            output_path = str(input_file.parent / f"{input_file.stem}_telugu{input_file.suffix}")
        
        print(f"Reading DOCX file: {input_path}")
        doc = Document(input_path)
        
        # Extract all text elements
        all_texts = []
        text_elements = []
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_texts.append(para.text)
                text_elements.append(('paragraph', para))
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            all_texts.append(para.text)
                            text_elements.append(('table_cell', para))
        
        # Headers and footers
        for section in doc.sections:
            for header in [section.header]:
                for para in header.paragraphs:
                    if para.text.strip():
                        all_texts.append(para.text)
                        text_elements.append(('header', para))
            
            for footer in [section.footer]:
                for para in footer.paragraphs:
                    if para.text.strip():
                        all_texts.append(para.text)
                        text_elements.append(('footer', para))
        
        if not all_texts:
            print("No text found in document")
            return output_path
        
        print(f"Translating {len(all_texts)} text segments...")
        
        # Translate all texts
        translations = translate_batch(all_texts, engine=self.engine, glossary=self.glossary)
        
        # Log translation
        if user_id:
            source_text = "\n\n".join(all_texts)
            translated_text = "\n\n".join(translations)
            log_translation(user_id, self.engine, source_text, translated_text, translation_id)
        
        # Replace text in document
        translation_idx = 0
        for element_type, element in text_elements:
            if translation_idx < len(translations):
                element.text = translations[translation_idx]
                translation_idx += 1
        
        # Save translated document with retry logic for file locking issues
        max_retries = 3
        for attempt in range(max_retries):
            try:
                doc.save(output_path)
                print(f"Translated document saved to: {output_path}")
                return output_path
            except PermissionError as e:
                if attempt < max_retries - 1:
                    import time
                    time.sleep(0.5)  # Wait 0.5 seconds before retry
                    print(f"Permission error saving file, retrying ({attempt + 1}/{max_retries})...")
                else:
                    print(f"Error saving document after {max_retries} attempts: {e}")
                    raise
            except Exception as e:
                print(f"Error saving document: {e}")
                raise

